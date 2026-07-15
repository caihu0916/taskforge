
# Copyright (c) 2024-2026 TaskForge Team
# SPDX-License-Identifier: BSL-1.1

"""Agent输出格式适配器 — 将结构化数据渲染为MD/DOCX/XLSX/PDF文件"""

from __future__ import annotations

import json
import os
from datetime import datetime
from pathlib import Path

import structlog

logger = structlog.get_logger(__name__)

# 基于项目根目录的绝对路径（不依赖os.getcwd()）
_PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent.parent  # src/engine/agent/output → 项目根
_DELIVERABLES_ROOT = str(_PROJECT_ROOT / "data" / "deliverables")
_TEMPLATES_ROOT = str(_PROJECT_ROOT / "data" / "templates")

# 文件大小限制
MAX_DELIVERABLE_CONTENT_CHARS = 5_000_000  # 500万字符 ≈ 5MB纯文本
MAX_XLSX_ROWS = 10_000


def get_deliverables_root() -> str:
    """获取交付物根目录"""
    os.makedirs(_DELIVERABLES_ROOT, exist_ok=True)
    return _DELIVERABLES_ROOT


def get_templates_root() -> str:
    """获取模板根目录"""
    os.makedirs(_TEMPLATES_ROOT, exist_ok=True)
    return _TEMPLATES_ROOT


def render_deliverable(
    format: str,
    content: str | dict,
    title: str = "TaskForge交付物",
    template_name: str | None = None,
    agent_name: str = "Agent",
    context: dict | None = None,
) -> dict:
    """将Agent输出渲染为目标格式文件

    Returns:
        {"file_path": str, "file_name": str, "file_type": str, "size_bytes": int}
    """
    # 内容大小校验
    content_size = len(content) if isinstance(content, str) else len(json.dumps(content, ensure_ascii=False))
    if content_size > MAX_DELIVERABLE_CONTENT_CHARS:
        logger.warning("content_truncated", original=content_size, limit=MAX_DELIVERABLE_CONTENT_CHARS)
        if isinstance(content, str):
            content = content[:MAX_DELIVERABLE_CONTENT_CHARS] + "\n\n[内容过长已截断]"
        else:
            content = {"warning": "内容过长已截断", "partial": str(content)[:MAX_DELIVERABLE_CONTENT_CHARS]}

    os.makedirs(_DELIVERABLES_ROOT, exist_ok=True)
    os.makedirs(_TEMPLATES_ROOT, exist_ok=True)

    date_str = datetime.now().strftime("%Y%m%d_%H%M%S")
    dir_name = f"{agent_name}_{date_str}"
    out_dir = os.path.join(_DELIVERABLES_ROOT, dir_name)
    os.makedirs(out_dir, exist_ok=True)

    if format == "md":
        return _render_md(content, title, out_dir)
    if format == "docx":
        return _render_docx(content, title, out_dir, template_name, context)
    if format == "xlsx":
        return _render_xlsx(content, title, out_dir, context)
    if format == "pdf":
        return _render_pdf(content, title, out_dir, context)
    logger.warning("unknown_format_fallback_md", format=format)
    return _render_md(content, title, out_dir)


def _render_md(content: str | dict, title: str, out_dir: str) -> dict:
    file_name = f"{_safe_filename(title)}.md"
    file_path = os.path.join(out_dir, file_name)

    if isinstance(content, dict):
        md_lines = [f"#{title}", ""]
        md_lines.extend(_dict_to_md(content))
        md_content = "\n".join(md_lines)
    else:
        md_content = f"#{title}\n\n{content}"

    with open(file_path, "w", encoding="utf-8") as f:
        f.write(md_content)

    return _file_result(file_path, file_name, "md")


def _render_docx(
    content: str | dict, title: str, out_dir: str, template_name: str | None, context: dict | None
) -> dict:
    from docxtpl import DocxTemplate as DocxTpl

    file_name = f"{_safe_filename(title)}.docx"
    file_path = os.path.join(out_dir, file_name)

    tpl_context = {"title": title, "date": datetime.now().strftime("%Y-%m-%d")}
    if isinstance(content, dict):
        tpl_context.update(content)
    elif isinstance(content, str):
        tpl_context["content"] = content
    if context:
        tpl_context.update(context)

    if template_name:
        tpl_path = os.path.join(_TEMPLATES_ROOT, template_name)
        if os.path.exists(tpl_path):
            try:
                doc = DocxTpl(tpl_path)
                doc.render(tpl_context)
                doc.save(file_path)
                logger.info("docx_rendered_with_template", template=template_name)
                return _file_result(file_path, file_name, "docx")
            except Exception as e:
                logger.warning("docx_template_failed_fallback", error=str(e), exc_info=True)

    _build_docx_from_scratch(tpl_context, file_path)
    return _file_result(file_path, file_name, "docx")


def _build_docx_from_scratch(context: dict, file_path: str) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    title = context.get("title", "TaskForge交付物")
    doc.add_heading(title, level=0)

    doc.add_paragraph(f"生成日期: {context.get('date', '')}").runs[0].font.size = Pt(9)

    content = context.get("content", "")
    if content:
        for para in content.split("\n\n"):
            stripped = para.strip()
            if not stripped:
                continue
            if stripped.startswith("#"):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("##"):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("###"):
                doc.add_heading(stripped[4:], level=3)
            else:
                doc.add_paragraph(stripped)

    skip_keys = {"title", "content", "date"}
    for key, val in context.items():
        if key in skip_keys:
            continue
        doc.add_heading(str(key).replace("_", " ").title(), level=2)
        if isinstance(val, (list, dict)):
            doc.add_paragraph(json.dumps(val, ensure_ascii=False, indent=2))
        else:
            doc.add_paragraph(str(val))

    doc.save(file_path)
    logger.info("docx_built_from_scratch")


def _render_xlsx(content: str | dict, title: str, out_dir: str, context: dict | None) -> dict:
    import openpyxl
    from openpyxl.styles import Border, Font, PatternFill, Side

    file_name = f"{_safe_filename(title)}.xlsx"
    file_path = os.path.join(out_dir, file_name)

    wb = openpyxl.Workbook()

    header_font = Font(bold=True, size=11, color="FFFFFF")
    header_fill = PatternFill(start_color="00F5FF", end_color="00F5FF", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    if isinstance(content, dict):
        for sheet_idx, (sheet_name, sheet_data) in enumerate(content.items()):
            if sheet_idx == 0:
                ws = wb.active
                ws.title = _safe_filename(sheet_name)[:31]
            else:
                ws = wb.create_sheet(_safe_filename(sheet_name)[:31])

            if isinstance(sheet_data, list) and sheet_data:
                # XLSX行数限制
                if len(sheet_data) > MAX_XLSX_ROWS:
                    logger.warning("xlsx_rows_truncated", original=len(sheet_data), limit=MAX_XLSX_ROWS)
                    truncated = sheet_data[:MAX_XLSX_ROWS]
                else:
                    truncated = sheet_data
                _write_table(ws, truncated, header_font, header_fill, thin_border)
            elif isinstance(sheet_data, dict):
                _write_kv_table(ws, sheet_data, header_font, header_fill, thin_border)
    elif isinstance(content, str):
        ws = wb.active
        ws.title = "数据"
        ws.append(["内容"])
        ws.append([content])

    wb.save(file_path)
    logger.info("xlsx_rendered")
    return _file_result(file_path, file_name, "xlsx")


def _render_pdf(content: str | dict, title: str, out_dir: str, context: dict | None) -> dict:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    # 注册中文字体 (Windows SimHei)
    _cn_font = "SimHei"
    _cn_font_paths = [
        os.path.join(
            os.environ.get("SYSTEMROOT", r"C:\Windows") if os.name == "nt" else r"C:\Windows", "Fonts", "simhei.ttf"
        ),
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",  # Linux fallback
    ]
    for fp in _cn_font_paths:
        if os.path.exists(fp):
            try:
                pdfmetrics.registerFont(TTFont(_cn_font, fp))
                break
            except Exception as exc:
                logger.debug("exception_handled", error=str(exc))
                _cn_font = "Helvetica"  # 降级
    else:
        _cn_font = "Helvetica"  # 无中文字体降级

    file_name = f"{_safe_filename(title)}.pdf"
    file_path = os.path.join(out_dir, file_name)

    doc = SimpleDocTemplate(file_path, pagesize=A4)
    styles = getSampleStyleSheet()
    # 覆盖默认字体为中文字体
    for style_name in ("Title", "Heading1", "Heading2", "Normal"):
        if style_name in styles.byName:
            styles.byName[style_name].fontName = _cn_font
    story = []

    story.append(Paragraph(title, styles["Title"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"生成日期: {datetime.now().strftime('%Y-%m-%d')}", styles["Normal"]))
    story.append(Spacer(1, 24))

    text = content if isinstance(content, str) else json.dumps(content, ensure_ascii=False, indent=2)
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 6))
        elif stripped.startswith("#"):
            story.append(Paragraph(stripped[2:], styles["Heading1"]))
        elif stripped.startswith("##"):
            story.append(Paragraph(stripped[3:], styles["Heading2"]))
        else:
            story.append(Paragraph(stripped, styles["Normal"]))

    doc.build(story)
    logger.info("pdf_rendered")
    return _file_result(file_path, file_name, "pdf")


def _safe_filename(name: str) -> str:
    """文件名安全化 — 保留中文、字母、数字、下划线、连字符"""
    import re

    # 移除路径分隔符和特殊字符，保留中文
    safe = re.sub(r'[\\/:*?"<>|\r\n\t\0]', "_", name)
    # 移除控制字符(\x00-\x1f)
    safe = re.sub(r"[\x00-\x1f]", "", safe)
    return safe.strip().strip(".")[:50] or "deliverable"


def _file_result(file_path: str, file_name: str, file_type: str) -> dict:
    return {
        "file_path": file_path,
        "file_name": file_name,
        "file_type": file_type,
        "size_bytes": os.path.getsize(file_path),
    }


def _write_table(ws, data: list, hfont, hfill, border) -> None:
    if not data:
        return
    if isinstance(data[0], dict):
        headers = list(data[0].keys())
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.font = hfont
            cell.fill = hfill
            cell.border = border
        for row_idx, item in enumerate(data, 2):
            for col, h in enumerate(headers, 1):
                cell = ws.cell(row=row_idx, column=col, value=item.get(h, ""))
                cell.border = border
    elif isinstance(data[0], list):
        for col, h in enumerate(data[0], 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.font = hfont
            cell.fill = hfill
            cell.border = border
        for row_idx, row in enumerate(data[1:], 2):
            for col, val in enumerate(row, 1):
                cell = ws.cell(row=row_idx, column=col, value=val)
                cell.border = border

    for col in ws.columns:
        max_len = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)


def _write_kv_table(ws, data: dict, hfont, hfill, border) -> None:
    ws.cell(row=1, column=1, value="字段").font = hfont
    ws.cell(row=1, column=1).fill = hfill
    ws.cell(row=1, column=1).border = border
    ws.cell(row=1, column=2, value="值").font = hfont
    ws.cell(row=1, column=2).fill = hfill
    ws.cell(row=1, column=2).border = border
    for row_idx, (k, v) in enumerate(data.items(), 2):
        ws.cell(row=row_idx, column=1, value=k).border = border
        ws.cell(row=row_idx, column=2, value=str(v) if not isinstance(v, str) else v).border = border
    ws.column_dimensions["A"].width = 25
    ws.column_dimensions["B"].width = 50


def _dict_to_md(data: dict, level: int = 2) -> list[str]:
    lines = []
    prefix = "#" * level
    for k, v in data.items():
        if isinstance(v, dict):
            lines.append(f"{prefix} {k}")
            lines.extend(_dict_to_md(v, level + 1))
        elif isinstance(v, list):
            lines.append(f"{prefix} {k}")
            for item in v:
                lines.append(f"- {item}")
        else:
            lines.append(f"**{k}**: {v}")
    return lines
